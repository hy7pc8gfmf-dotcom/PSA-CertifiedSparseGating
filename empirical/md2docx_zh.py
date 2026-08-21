# -*- coding: utf-8 -*-
"""md → docx 转换（中文学术排版）：标题/表格/代码块/列表/引用/粗体/行内代码"""
import re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CN_BODY = '宋体'; CN_HEAD = '黑体'; CN_MONO = '等线'; EN_BODY = 'Times New Roman'; EN_MONO = 'Consolas'

def _set(run, size=12, bold=False, italic=False, mono=False, cn=CN_BODY, en=EN_BODY):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = en if not mono else EN_MONO
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), CN_MONO if mono else cn)

def add_inline(par, text, size=12, base_bold=False, italic_all=False):
    tokens = re.split(r'(\*\*.*?\*\*|`[^`]*`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = par.add_run(tok[2:-2]); _set(r, size, bold=True, italic=italic_all)
        elif tok.startswith('`') and tok.endswith('`'):
            r = par.add_run(tok[1:-1]); _set(r, size - 1, mono=True)
        else:
            for p in re.split(r'(\*[^*\s][^*]*?\*)', tok):
                if len(p) > 1 and p.startswith('*') and p.endswith('*') and not p.startswith('**'):
                    r = par.add_run(p[1:-1]); _set(r, size, italic=True)
                else:
                    r = par.add_run(p); _set(r, size, bold=base_bold, italic=italic_all)

def flush_table(doc, buf):
    rows = []
    for r in buf:
        s = r.strip().strip('|')
        if s:
            rows.append([c.strip() for c in s.split('|')])
    rows = [r for r in rows if not all(re.fullmatch(r':?-{2,}:?', c) for c in r)]
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = t.cell(ri, ci)
            par = cell.paragraphs[0]
            add_inline(par, row[ci] if ci < len(row) else '', size=10.5)
            if ri == 0:
                for r in par.runs:
                    r.bold = True
    doc.add_paragraph()

def md_to_docx(md_path, docx_path):
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(1)
        sec.top_margin = sec.bottom_margin = Inches(1)
    st = doc.styles['Normal']
    st.font.name = EN_BODY; st.font.size = Pt(12)
    st._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), CN_BODY)

    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    i, in_code, code_buf, table_buf = 0, False, [], []

    def emit_heading(level, text):
        par = doc.add_heading(level=level)
        for r in list(par.runs):
            r.text = ''
        add_inline(par, text, size={1: 16, 2: 14, 3: 12.5}.get(level, 12))
        for r in par.runs:
            r.bold = True
            r.font.name = EN_BODY
            r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), CN_HEAD)
        if level == 1:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if s.startswith('```'):
            if in_code:
                par = doc.add_paragraph()
                for j, cl in enumerate(code_buf):
                    if j:
                        par.add_run('\n')
                    r = par.add_run(cl); _set(r, 10, mono=True)
                in_code = False; code_buf = []
            else:
                flush_table(doc, table_buf); table_buf = []
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue
        if not s:
            flush_table(doc, table_buf); table_buf = []
            i += 1; continue
        if s.startswith('|'):
            table_buf.append(line); i += 1; continue
        flush_table(doc, table_buf); table_buf = []
        if re.fullmatch(r'-{3,}|={3,}|\*{3,}', s):
            i += 1; continue
        m = re.match(r'^(#{1,6})\s+(.*)$', s)
        if m:
            emit_heading(len(m.group(1)), m.group(2)); i += 1; continue
        if s.startswith('>'):
            body = s.lstrip('>').strip()
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.3)
            add_inline(par, body, size=10.5, italic_all=True)
            for r in par.runs:
                r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1; continue
        m = re.match(r'^\s*[-*]\s+(.*)$', s)
        if m:
            par = doc.add_paragraph(style='List Bullet')
            add_inline(par, m.group(1)); i += 1; continue
        m = re.match(r'^\s*(\d+)[.、]\s+(.*)$', s)
        if m:
            par = doc.add_paragraph(style='List Number')
            add_inline(par, m.group(2)); i += 1; continue
        par = doc.add_paragraph()
        par.paragraph_format.first_line_indent = Pt(24)
        par.paragraph_format.line_spacing = 1.4
        add_inline(par, s)
        i += 1

    flush_table(doc, table_buf)
    doc.save(docx_path)
    print(f'OK: {docx_path}')

if __name__ == '__main__':
    md_to_docx(sys.argv[1], sys.argv[2])
